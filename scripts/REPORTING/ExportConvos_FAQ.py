import sys
import requests
import os
import json
from datetime import datetime
import openpyxl  # Přidejte tento import na začátek souboru
from openpyxl.styles import Font, Alignment
import re
from openpyxl.chart import PieChart, Reference
from openpyxl.styles import PatternFill
import random
import matplotlib.pyplot as plt
from io import BytesIO
from openpyxl.drawing.image import Image
import numpy as np
import matplotlib.colors as mcolors
from collections import Counter
from anthropic import Anthropic  # Add this import
from openai import OpenAI  # Add this import for OpenRouter
import logging
from docx import Document  # Add this import
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Funkce pro načtení konfigurace ze souboru
def load_config(config_file):
    logging.info(f"Loading configuration from {config_file}")
    config = {}
    try:
        with open(config_file, 'r') as file:
            for line in file:
                key, value = line.strip().split('=', 1)
                if key == 'CATEGORIES':
                    config[key] = value.strip('[]').split(',')
                elif key == 'OPENROUTER_MODELS':
                    # Parse JSON array for OpenRouter models
                    config[key] = json.loads(value.strip())
                else:
                    config[key] = value.strip()
        logging.info("Configuration loaded successfully")
        logging.debug(f"Config contents: {config}")
        return config
    except Exception as e:
        logging.error(f"Error loading config: {str(e)}")
        raise

# Function to initialize LLM client based on provider
def initialize_llm_client(config):
    """Initialize the appropriate LLM client based on configuration"""
    provider = config.get('LLM_PROVIDER', 'anthropic').lower()
    
    if provider == 'openrouter':
        logging.info("Initializing OpenRouter client")
        api_key = config.get('OPENROUTER_API_KEY')
        if not api_key:
            raise ValueError("OPENROUTER_API_KEY is required when using OpenRouter provider")
        
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key
        )
        
        models = config.get('OPENROUTER_MODELS', ['anthropic/claude-3.5-sonnet'])
        logging.info(f"OpenRouter client initialized with models: {models}")
        return client, 'openrouter', models
    
    elif provider == 'anthropic':
        logging.info("Initializing Anthropic client")
        api_key = config.get('ANTHROPIC_API_KEY')
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY is required when using Anthropic provider")
        
        client = Anthropic(api_key=api_key)
        models = ['claude-sonnet-4-20250514']  # Default Anthropic model
        logging.info(f"Anthropic client initialized with model: {models[0]}")
        return client, 'anthropic', models
    
    else:
        raise ValueError(f"Unsupported LLM provider: {provider}")

# Set up logging configuration
def setup_logging(output_directory):
    log_filepath = os.path.join(output_directory, 'execution.log')
    
    # First remove any existing handlers to avoid duplicates
    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    
    # Configure logging with immediate output
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    
    file_handler = logging.FileHandler(log_filepath, encoding='utf-8', mode='w')
    file_handler.setLevel(logging.INFO)
    
    # Create formatter and add it to the handlers
    formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
    console_handler.setFormatter(formatter)
    file_handler.setFormatter(formatter)
    
    # Get the root logger and add handlers
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    logger.addHandler(console_handler)
    logger.addHandler(file_handler)
    
    # Force immediate flush
    console_handler.flush()
    file_handler.flush()
    
    # Test logging is working
    logging.info("="*80)
    logging.info("LOGGING INITIALIZED SUCCESSFULLY")
    logging.info(f"Log file created at: {log_filepath}")
    logging.info("="*80)
    
    return log_filepath

def main(config_file):
    try:
        # Create output directory first
        config = load_config(config_file)
        
        global PROJECT_NAME, START_DATE, END_DATE, OUTPUT_DIRECTORY, CATEGORIES, ANTHROPIC_API_KEY
        PROJECT_NAME = config['PROJECT_NAME']
        START_DATE = config['START_DATE']
        END_DATE = config['END_DATE']
        OUTPUT_DIRECTORY = f"{PROJECT_NAME}_{START_DATE}_to_{END_DATE}"
        
        # Create directory and set up logging immediately
        create_output_directory()
        setup_logging(OUTPUT_DIRECTORY)  # Move this up, right after directory creation
        
        logging.info(f"Starting script with config file: {config_file}")
        
        # Rest of your configuration
        auth_token = config['AUTH_TOKEN']
        project_id = config['PROJECT_ID']
        CATEGORIES = config['CATEGORIES']
        generate_faq = config.get('GenerateFAQ', 'true').lower() == 'true'  # New parameter with default value
        
        # Initialize LLM client based on provider
        llm_client, llm_provider, llm_models = initialize_llm_client(config)
        
        base_url = "https://api.voiceflow.com/v2/transcripts"
        headers = {
            "Authorization": auth_token,
            "accept": "application/json"
        }

        # Create output directory first
        create_output_directory()
        
        # Set up logging immediately after creating output directory
        log_filepath = setup_logging(OUTPUT_DIRECTORY)
        logging.info(f"Starting script execution for project: {PROJECT_NAME}")
        logging.info(f"Date range: {START_DATE} to {END_DATE}")
        
        transcript_ids = get_transcript_ids(base_url, headers, project_id)
        
        total_human_count = 0
        
        for transcript_id in transcript_ids:
            dialog = get_transcript_dialog(base_url, headers, project_id, transcript_id)
            messages = extract_messages(dialog)
            save_transcript_to_txt(transcript_id, messages)
            
            for message in messages:
                if message['role'] == 'HUMAN':
                    total_human_count += 1
        
        category_counts = count_category_occurrences()
        create_excel_report(total_human_count, category_counts)
        
        if generate_faq:
            logging.info("FAQ generation is enabled, proceeding with conversation analysis...")
            print(f"Analyzing conversations with LLM ({llm_provider})...")
            faq_analysis = analyze_conversations_with_llm(OUTPUT_DIRECTORY, llm_client, llm_provider, llm_models, total_human_count)
            save_faq_to_word(faq_analysis)
        else:
            logging.info("FAQ generation is disabled, skipping conversation analysis...")
        
        print("Zpracování dokončeno.")
        
        logging.info("Script execution completed successfully")
    except Exception as e:
        logging.error(f"Critical error in main execution: {str(e)}", exc_info=True)
        raise
    finally:
        logging.info("="*80)
        logging.info("SCRIPT EXECUTION ENDED")
        logging.info("="*80)

def create_output_directory():
    logging.info(f"Creating output directory: {OUTPUT_DIRECTORY}")
    try:
        if not os.path.exists(OUTPUT_DIRECTORY):
            os.makedirs(OUTPUT_DIRECTORY)
            logging.info(f"Created directory: {OUTPUT_DIRECTORY}")
        else:
            logging.info(f"Directory already exists: {OUTPUT_DIRECTORY}")
    except Exception as e:
        logging.error(f"Error creating output directory: {str(e)}")
        raise

def get_transcript_ids(base_url, headers, project_id):
    logging.info(f"Fetching transcript IDs for project {project_id}")
    try:
        url = f"{base_url}/{project_id}"
        params = {
            "startDate": START_DATE,
            "endDate": END_DATE
        }
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        transcript_ids = [transcript["_id"] for transcript in response.json()]
        logging.info(f"Retrieved {len(transcript_ids)} transcript IDs")
        return transcript_ids
    except Exception as e:
        logging.error(f"Error fetching transcript IDs: {str(e)}")
        raise

def get_transcript_dialog(base_url, headers, project_id, transcript_id):
    logging.info(f"Fetching dialog for transcript {transcript_id}")
    try:
        url = f"{base_url}/{project_id}/{transcript_id}"
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        dialog = response.json()
        logging.info(f"Successfully retrieved dialog for transcript {transcript_id} - {len(dialog)} messages")
        return dialog
    except Exception as e:
        logging.error(f"Error fetching transcript dialog {transcript_id}: {str(e)}")
        raise

def extract_messages(dialog):
    messages = []
    for turn in dialog:
        if turn["type"] == "debug" and "payload" in turn and "payload" in turn["payload"]:
            debug_message = turn["payload"]["payload"].get("message", "")
            if "CategoryFilter" in debug_message:
                messages.append({
                    "role": "DEBUG",
                    "content": debug_message,
                    "timestamp": turn.get("startTime", "")
                })
        elif turn["type"] == "request":
            if "payload" in turn and "query" in turn["payload"].get("payload", {}):
                messages.append({
                    "role": "HUMAN",
                    "content": turn["payload"]["payload"]["query"],
                    "timestamp": turn.get("startTime", "")
                })
        elif turn["type"] == "text" and "payload" in turn and "message" in turn["payload"].get("payload", {}):
            messages.append({
                "role": "BOT",
                "content": turn["payload"]["payload"]["message"],
                "timestamp": turn.get("startTime", "")
            })
    return messages

def save_transcript_to_txt(transcript_id, messages):
    filename = os.path.join(OUTPUT_DIRECTORY, f"transcript_{transcript_id}.txt")
    logging.info(f"Saving transcript {transcript_id} to file...")
    sys.stdout.flush()  # Force immediate console output
    
    try:
        message_count = len(messages)
        with open(filename, 'w', encoding='utf-8') as file:
            for idx, message in enumerate(messages, 1):
                if message['role'] == 'DEBUG':
                    file.write(f"DEBUG: {message['content']}\n")
                    file.write(f"TIMESTAMP: {message['timestamp']}\n")
                else:
                    file.write(f"{message['role']}: {message['content']}\n")
                    file.write(f"TIMESTAMP: {message['timestamp']}\n")
                file.write("----------\n")
        logging.info(f"✓ Successfully saved transcript {transcript_id} ({message_count} messages)")
        sys.stdout.flush()  # Force immediate console output
    except Exception as e:
        logging.error(f"Error saving transcript {transcript_id}: {str(e)}")
        sys.stdout.flush()  # Force immediate console output
        raise

def count_human_occurrences():
    human_count = 0
    for filename in os.listdir(OUTPUT_DIRECTORY):
        if filename.startswith("transcript_") and filename.endswith(".txt"):
            file_path = os.path.join(OUTPUT_DIRECTORY, filename)
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                human_count += content.count("HUMAN:")
    return human_count

def count_category_occurrences():
    category_counts = {category: 0 for category in CATEGORIES}
    
    for filename in os.listdir(OUTPUT_DIRECTORY):
        if filename.startswith("transcript_") and filename.endswith(".txt"):
            file_path = os.path.join(OUTPUT_DIRECTORY, filename)
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                for category in CATEGORIES:
                    exact_match = f'\\"{category}\\"'
                    category_counts[category] += content.count(exact_match)
    
    return category_counts

def create_pie_chart(category_counts):
    labels = list(category_counts.keys())
    sizes = list(category_counts.values())
    
    plt.figure(figsize=(10, 8))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    plt.axis('equal')
    plt.title('Rozložení kategorií')
    
    # Uložení grafu jako obrázek
    img_path = os.path.join(OUTPUT_DIRECTORY, 'category_distribution.png')
    plt.savefig(img_path)
    plt.close()
    
    return img_path

def create_custom_donut_chart(category_counts):
    logging.info("Creating custom donut chart")
    try:
        # Seřazení kategorií sestupně podle počtu
        sorted_categories = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)
        labels = [cat for cat, count in sorted_categories if count > 0]
        sizes = [count for _, count in sorted_categories if count > 0]

        # Vytvoření širšího spektra barev duhy v opačném pořadí
        num_colors = len(labels)
        rainbow_colors = plt.cm.rainbow(np.linspace(1, 0, num_colors))

        # Změna velikosti figury na čtvercový tvar
        fig, ax = plt.subplots(figsize=(20, 20), subplot_kw=dict(aspect="equal"))  

        wedges, texts, autotexts = ax.pie(sizes, wedgeprops=dict(width=0.5), startangle=-40,
                                          colors=rainbow_colors, autopct='%1.1f%%', pctdistance=0.85)

        bbox_props = dict(boxstyle="round,pad=0.3", fc="w", ec="k", lw=0.72)
        kw = dict(arrowprops=dict(arrowstyle="-", connectionstyle="angle,angleA=0,angleB=90,rad=10"),
                  bbox=bbox_props, zorder=0, va="center")

        # Úprava umístění popisků
        for i, p in enumerate(wedges):
            ang = (p.theta2 - p.theta1) / 2. + p.theta1
            y = np.sin(np.deg2rad(ang))
            x = np.cos(np.deg2rad(ang))
            horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
            connectionstyle = f"angle,angleA=0,angleB={ang}"
            kw["arrowprops"].update({"connectionstyle": connectionstyle})
            ax.annotate(labels[i], xy=(x, y), xytext=(2.2*np.sign(x), 2.2*y),
                        horizontalalignment=horizontalalignment, fontsize=14, **kw)

        plt.title(f"Rozložení kategorií\n{START_DATE} - {END_DATE}", fontsize=24, y=1.05)

        for autotext in autotexts:
            autotext.set_visible(False)

        # Odstranění os pro čistší vzhled
        ax.set_axis_off()

        # Uložení grafu
        img_path = os.path.join(OUTPUT_DIRECTORY, 'custom_category_distribution.png')
        plt.savefig(img_path, bbox_inches='tight', dpi=300)
        plt.close()
        logging.info(f"Donut chart saved to {img_path}")
        return img_path
    except Exception as e:
        logging.error(f"Error creating donut chart: {str(e)}")
        raise

def create_excel_report(ai_responses_count, category_counts):
    logging.info("\n" + "="*50)
    logging.info("CREATING EXCEL REPORT")
    logging.info(f"Total AI responses: {ai_responses_count}")
    logging.info(f"Categories to process: {len(category_counts)}")
    
    try:
        # Vytvoření nového Excel sešitu
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Report"

        # Nastavení nadpisu
        ws['A1'] = f"Report kategorií za období {START_DATE} - {END_DATE}"
        ws['A1'].font = openpyxl.styles.Font(size=16, bold=True)

        # Přidání počtu AI odpovědí
        ws['A2'] = "Počet AI odpovědí utracených za dané období:"
        ws['B2'] = ai_responses_count

        # Zvýraznění počtu AI odpovědí
        ws['A2'].font = openpyxl.styles.Font(size=14, bold=True, color="FF0000")
        ws['B2'].font = openpyxl.styles.Font(size=14, bold=True, color="FF0000")
        ws['A2'].fill = openpyxl.styles.PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        ws['B2'].fill = openpyxl.styles.PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

        # Přidání informace o celkovém počtu kategorizac
        total_categorizations = sum(category_counts.values())
        ws['A4'] = f"Celkový počet přiřazení / kategorizací dotazů: {total_categorizations}"
        ws['A5'] = "Poznámka: 1 dotaz může být přiřazen do více kategorií, proto se celkový počet AI odpovědí nerovná počtu kategorizací."

        # Přidání hlavičky tabulky
        ws['A7'] = "Kategorie"
        ws['B7'] = "Počet"
        for cell in ws['A7:B7'][0]:
            cell.font = openpyxl.styles.Font(bold=True)

        # Seřazení kategorií sestupně podle počtu
        sorted_categories = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)

        # Přidání dat kategorií
        for row, (category, count) in enumerate(sorted_categories, start=8):
            ws.cell(row=row, column=1, value=category)
            ws.cell(row=row, column=2, value=count)

        # Ohraničení tabulky
        max_row = ws.max_row
        for row in ws[f'A7:B{max_row}']:
            for cell in row:
                cell.border = openpyxl.styles.Border(left=openpyxl.styles.Side(style='thin'),
                                                     right=openpyxl.styles.Side(style='thin'),
                                                     top=openpyxl.styles.Side(style='thin'),
                                                     bottom=openpyxl.styles.Side(style='thin'))

        # Nastavení šířky prvního sloupce
        ws.column_dimensions['A'].width = 150

        # Vytvoření vlastního grafu
        img_path = create_custom_donut_chart(category_counts)

        # Vložení obrázku do Excel souboru pod tabulkou
        img = Image(img_path)
        img.width = 1000
        img.height = 800
        ws.add_image(img, f'A{max_row + 2}')

        # Uložení Excel souboru
        output_filename = f'{PROJECT_NAME}_report_{START_DATE}_to_{END_DATE}.xlsx'
        output_path = os.path.join(OUTPUT_DIRECTORY, output_filename)
        wb.save(output_path)
        logging.info(f"Excel report successfully created: {output_path}")

        print(f"Nový Excel report byl vytvořen a uložen: {output_path}")
        
        logging.info("\nCategory distribution:")
        for category, count in sorted_categories:
            logging.info(f"  • {category}: {count}")
        
        logging.info(f"\n✓ Excel report successfully created: {output_path}")
        logging.info("="*50)
    except Exception as e:
        logging.error(f"Error creating Excel report: {str(e)}")
        raise

def print_summary(human_count, category_counts):
    print("\nSouhrn zpracování:")
    print(f"1) Počet HUMAN výskytů: {human_count}")
    print("\n2) Počty výskytů kategorií:")
    for category, count in category_counts.items():
        print(f"   {category}: {count}")

def analyze_conversations_with_llm(transcripts_directory, client, provider, models, total_human_count):
    logging.info(f"Starting conversation analysis with {provider}")
    logging.info(f"Available models: {models}")
    logging.info(f"Total human messages from Excel: {total_human_count}")
    
    qa_pairs = []
    transcript_count = 0
    
    logging.info(f"Scanning directory: {transcripts_directory}")
    for filename in os.listdir(transcripts_directory):
        if filename.startswith("transcript_") and filename.endswith(".txt"):
            transcript_count += 1
            file_path = os.path.join(transcripts_directory, filename)
            logging.info(f"Processing transcript file: {filename}")
            
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    # Split by delimiter
                    messages = content.split('----------')
                    
                    current_question = None
                    current_question_text = []
                    qa_count = 0
                    
                    for message in messages:
                        message = message.strip()
                        if not message:
                            continue
                            
                        # Log message for debugging
                        logging.debug(f"Processing message:\n{message[:200]}...")
                        
                        if message.startswith("HUMAN: "):
                            # Start collecting new question
                            current_question_text = [message[7:].strip()]
                        elif message.startswith("BOT: ") and current_question_text:
                            # Get full answer text
                            answer = message[5:].strip()
                            
                            # Combine multi-line question if any
                            question = "\n".join(current_question_text)
                            
                            # Add validation
                            if len(question) > 0 and len(answer) > 0:
                                qa_pairs.append((question, answer))
                                qa_count += 1
                                
                                # Log extracted pair for debugging
                                logging.debug(f"Extracted Q&A pair {qa_count}:")
                                logging.debug(f"Q: {question[:100]}...")
                                logging.debug(f"A: {answer[:100]}...")
                            
                            current_question_text = []
                        elif current_question_text:
                            # Append to multi-line question
                            current_question_text.append(message)
                    
                    logging.info(f"Extracted {qa_count} Q&A pairs from {filename}")
            
            except Exception as e:
                logging.error(f"Error processing file {filename}: {str(e)}")
                continue

    logging.info(f"Total transcripts processed: {transcript_count}")
    logging.info(f"Total Q&A pairs collected: {len(qa_pairs)}")

    batch_size = 100
    return analyze_conversations_in_batches(client, provider, models, qa_pairs, total_human_count, batch_size=batch_size)

def make_llm_api_call(client, provider, models, prompt):
    """Make API call to the appropriate LLM provider"""
    if provider == 'anthropic':
        try:
            message = client.messages.create(
                model=models[0],
                max_tokens=8000,
                temperature=0,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            return message.content[0].text.strip()
        except Exception as e:
            logging.error(f"Anthropic API call failed: {str(e)}")
            raise
    
    elif provider == 'openrouter':
        try:
            # Prepare the request parameters
            request_params = {
                "model": models[0],  # Primary model
                "max_tokens": 8000,
                "temperature": 0,
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            }
            
            # Add all models (including primary) in extra_body for fallback support
            if len(models) > 1:
                request_params["extra_body"] = {
                    "models": models  # All models for fallback routing
                }
            
            completion = client.chat.completions.create(**request_params)
            return completion.choices[0].message.content.strip()
        except Exception as e:
            logging.error(f"OpenRouter API call failed: {str(e)}")
            raise
    
    else:
        raise ValueError(f"Unsupported provider: {provider}")

def save_faq_to_word(analysis_text):
    doc = Document()
    
    # Convert dates to Czech format
    start_date_obj = datetime.strptime(START_DATE, '%Y-%m-%d')
    end_date_obj = datetime.strptime(END_DATE, '%Y-%m-%d')
    start_date_cz = start_date_obj.strftime('%d.%m.%Y')
    end_date_cz = end_date_obj.strftime('%d.%m.%Y')
    
    # Nastavení stylu nadpisu
    title = doc.add_heading('', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title text with bold project name
    title_text = title.add_run(f"ANALÝZA NEJČASTĚJI KLADENÝCH DOTAZŮ ZA OBDOBÍ {start_date_cz} - {end_date_cz} PRO PROJEKT ")
    project_name = title.add_run(PROJECT_NAME)
    project_name.bold = True
    
    doc.add_paragraph()
    
    # Rozdělení analýzy na řádky
    lines = analysis_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Zpracování markdown nadpisů
        if line.startswith('#'):
            # Určení úrovně nadpisu podle počtu #
            level = len(line) - len(line.lstrip('#'))
            # Odstranění # a případného | ze začátku textu
            heading_text = line.lstrip('#').strip()
            if '|' in heading_text:
                heading_text = heading_text.split('|', 1)[1].strip()
            
            # Přidání nadpisu odpovídající úrovně (0-8)
            heading = doc.add_heading(heading_text, level)
            continue
        
        # Zpracování odrážek (bullet points)
        if line.startswith('*'):
            text = line[1:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            
            # Zpracování tučného textu mezi **
            if '**' in text:
                p.clear()
                parts = text.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Lichá část (mezi ** **)
                        p.add_run(part).bold = True
                    else:  # Sudá část (normální text)
                        p.add_run(part)
        
        # Běžný text
        else:
            p = doc.add_paragraph(line)
    
    # Nastavení stylu dokumentu
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(12)
        for run in paragraph.runs:
            run.font.size = Pt(11)
    
    # Uložení dokumentu
    output_filename = f'{PROJECT_NAME}_Často_kladené_dotazy_{START_DATE}_to_{END_DATE}.docx'
    output_path = os.path.join(OUTPUT_DIRECTORY, output_filename)
    doc.save(output_path)
    logging.info(f"FAQ analýza uložena do Word dokumentu: {output_path}")

def analyze_conversations_in_batches(client, provider, models, qa_pairs, total_human_count, batch_size=100):
    """Analyze conversations in batches to handle large datasets"""
    logging.info("="*80)
    logging.info("STARTING BATCH ANALYSIS PROCESS")
    logging.info(f"Total Q&A pairs to analyze: {len(qa_pairs)}")
    logging.info(f"Total human messages (from Excel): {total_human_count}")
    logging.info(f"Batch size: {batch_size}")
    logging.info("="*80)
    
    batch_analyses = []
    total_batches = (len(qa_pairs) + batch_size - 1) // batch_size
    
    for batch_num in range(total_batches):
        logging.info("\n" + "="*50)
        logging.info(f"PROCESSING BATCH {batch_num + 1}/{total_batches}")
        
        # Batch preparation
        start_idx = batch_num * batch_size
        end_idx = min((batch_num + 1) * batch_size, len(qa_pairs))
        current_batch = qa_pairs[start_idx:end_idx]
        
        logging.info(f"Batch size: {len(current_batch)} pairs")
        logging.info(f"Processing pairs {start_idx} to {end_idx}")
        
        # Validate qa_pairs before formatting
        for i, (q, a) in enumerate(current_batch):
            if not q or not a:
                logging.warning(f"Empty Q&A pair found at index {i}")
                continue
                
            if len(q) > 1000 or len(a) > 1000:
                logging.warning(f"Unusually long Q&A pair at index {i}:")
                logging.warning(f"Q length: {len(q)}, A length: {len(a)}")
                logging.warning(f"Q preview: {q[:100]}...")
                logging.warning(f"A preview: {a[:100]}...")
        
        # Format qa_text with additional validation
        qa_text = "\n\n".join([
            f"Q: {q}\nA: {a}" 
            for q, a in current_batch 
            if q and a and len(q) < 1000 and len(a) < 1000
        ])
        
        # Log sample of formatted text
        logging.debug(f"Formatted qa_text sample:\n{qa_text[:500]}...")
        
        # Create batch text and estimate tokens
        estimated_tokens = len(qa_text) // 4  # Rough estimation
        logging.info(f"Estimated tokens for this batch: {estimated_tokens}")
        
        batch_prompt = f"""# TVÁ IDENTITA: Jsi expertním datovým analytikem a legálním poradcem, který extrahuje NEJČASTĚJŠÍ páry dotazů od uživatele ("HUMAN") a NEJČASTĚJŠÍ odpovedí od umělé inteligence ("BOT")
        
        # INFORMACE PRO ROZEZNÁNÍ DOTAZU A ODPOVĚDI:
        1) **Dotaz:** Uživatelský dotaz rozeznáš podle klíčového slova "HUMAN"
        2) **Odpověď:** AI Odpověď na dotaz uživatele rozeznáš podle klíčového slova "BOT"
        3) **Zpětná vazba:** Analyzuj následnou reakci uživatele (pokud existuje) pro vyhodnocení spokojenosti
        
        # TVŮJ ULTIMÁTNÍ ÚKOL:
        1. Analyzuj pečlivě "TRANSKRIPCE KONVERZACE" a identifikuj skutečně NEJČASTĚJŠÍ a NEJDŮLEŽITĚJŠÍ témata dotazů uvnitř všech HUMAN & BOT párů dotazů a odpovědí jejich variace
        2. **SLEDUJ ČETNOSTI A SESKUPUJ INTELIGENTNĚ:** Pro každé téma spočítej PŘESNÝ počet výskytů významově podobných dotazů a seskupuj i drobné variace (např. "dej mi číslo na Frontze" + "dej mi číslo na pana Frontze" = 2 výskyty téhož dotazu)
        3. Vyber a seřaď PŘESNĚ 15 NEJRELEVANTNĚJŠÍCH témat podle četnosti jejich výskytů
        4. **IDENTIFIKUJ CIZÍ JAZYKY:** Rozpoznej všechny dotazy v jiných jazycích než češtině a pro každý jazyk identifikuj 3 nejčastější dotazy
        5. DŮLEŽITÉ: Poskytni KOMPLETNÍ analýzu všech 15 témat najednou, bez rozdělování do částí
        6. Pro každé téma analyzuj zpětnou vazbu uživatelů
        7. **DŮLEŽITÉ - POUZE RELEVANTNÍ DOTAZY:** Zaměř se výhradně na relevantní dotazy, které vyžadovaly informace k zodpovězení

        -----------------------

        # VYŽADOVANÝ FORMÁT VÝSTUPU: Pro každé z 15 témat použij tento formát:

        ## **[POŘADOVÉ ČÍSLO]** | [SOUHRNNÝ NÁZEV TÉMATU] ([CELKOVÝ POČET VÝSKYTŮ]x)
        ### ([SEZNAM NEJČASTĚJŠÍCH VARIANT s četnostmi - např. "kdy byla uzavřena Duchcovská?" (2x), "uzavírka Duchcovské ulice" (1x)])
        * **SUMARIZACE ODPOVĚĎÍ CHATBOTA:** [SOUHRNNÉ SUMARIZACE JAKÝM ZPŮSOBEM CHATBOT DOKÁZAL ODPOVĚDĚT - včetně kvality strukturování (URL, čísla, data, faktické informace)]

        [VLOŽIT 1 PRÁZDNÝ ŘÁDEK A POKRAČOVAT AŽ DO TÉMATU 15.]

        # SOUHRNNÁ STATISTIKA:
        
        ## CELKOVÁ PŘESNOST ANALÝZY
        * **PRŮMĚRNÁ MÍRA PŘESNOSTI:** [VYPOČÍTEJ PRŮMĚR VŠECH MĚR PŘESNOSTI Z VÝŠE UVEDENÝCH 15 TÉMAT]%
        * **ZDŮVODNĚNÍ:** [STRUČNÉ VYSVĚTLENÍ PRŮMĚRNÉ MÍRY PŘESNOSTI]

        ## CELKOVÁ ZPĚTNÁ VAZBA
        * **CELKOVÁ SPOKOJENOST UŽIVATELŮ:** [PRŮMĚR VŠECH POZITIVNÍCH ZPĚTNÝCH VAZEB]%
        * **ANALÝZA TRENDU:** [IDENTIFIKACE HLAVNÍCH DŮVODŮ SPOKOJENOSTI/NESPOKOJENOSTI]

        ## JAZYKOVÁ DISTRIBUCE: 
        * **ČEŠTINA:** [PROCENTO]%
        * **OSTATNÍ JAZYKY:** Pokud byly identifikovány dotazy v jiných jazycích, pro každý jazyk uveď:
          - **[NÁZEV JAZYKA]:** [PROCENTO]%
          - **TOP 3 NEJČASTĚJŠÍ DOTAZY:** 1) "[dotaz 1]", 2) "[dotaz 2]", 3) "[dotaz 3]"

        ## ANALÝZA NEJVYTĚŽOVANĚJŠÍCH HODIN
        * **NEJAKTIVNĚJŠÍ HODINY:** [IDENTIFIKUJ 3-5 KONKRÉTNÍCH HODIN V PRŮBĚHU DNE, KDY BYLO NEJVÍCE DOTAZŮ]
        * **ANALÝZA TRENDŮ:** [VYSVĚTLENÍ PROČ JSOU TYTO HODINY NEJVYTĚŽOVANĚJŠÍ]

        ## ODHADOVANÁ UŠETŘENÁ PRÁCA A FINANCE
        * **CELKOVÝ POČET DOTAZŮ:** {total_human_count} dotazů (ze statistik v Excel reportu)
        * **ODHADOVANÉ UŠETŘENÉ HODINY:** {total_human_count * 0.25} hodin (předpoklad 15 minut na dotaz)
        * **ODHADOVANÉ UŠETŘENÉ FINANCE:** {total_human_count * 0.25 * 200} Kč (při sazbě 200 Kč/hodina pro úředníka)
        * **ROČNÍ PROJEKCE:** [EXTRAPOLACE NA CELÝ ROK PODLE SOUČASNÉHO TRENDU]
        -----------------------

        
        # CELKOVÁ OMEZENÍ A POŽADAVKY:
        - Musíš analyzovat a vrátit VŠECH 15 témat najednou
        - **POVINNÉ ČETNOSTI:** Pro každé téma a variantu MUSÍŠ uvést přesný počet výskytů ve formátu "(Xx)"
        - **INTELIGENTNÍ SESKUPOVÁNÍ:** Seskupuj významově totožné dotazy i s drobnými variacemi (např. "dej mi číslo na Frontze" + "dej mi číslo na pana Frontze" = "dej mi číslo na (pana) Frontze" (2x))
        - **POVINNÉ CIZÍ JAZYKY:** Pokud existují dotazy v jiných jazycích než češtině, MUSÍŠ je identifikovat a uvést top 3 pro každý jazyk (BEZ ČÍSEL VÝSKYTŮ - POUZE PROCENTA)
        - **POUZE RELEVANTNÍ DOTAZY:** Zaměř se výhradně na relevantní dotazy, které vyžadovaly relevantní informace k zodpovězení
        - Nepoužívej žádné poznámky o pokračování nebo rozdělování odpovědi
        - Drž se striktně daného formátu pro každé téma
        - Vynechej na začátku tvého výstupu jakékoliv poznámky, či sebe vysvětlování => ihned začni s generováním jednotlivých dotazů a odpovědí
        - Pro každé téma MUSÍŠ analyzovat zpětnou vazbu uživatelů
        - Na konci MUSÍŠ uvést souhrnnou statistiku s průměrnou mírou přesnosti a celkovou zpětnou vazbou
        - **POVINNÉ NOVÉ SEKCE:** MUSÍŠ přidat analýzu nejvytěžovanějších hodin a odhadovanou ušetřenou práci a finance

        # "TRANSKRIPCE KONVERZACE": 
        {qa_text}

        """

        try:
            logging.info(f"Making {provider} API call...")
            start_time = datetime.now()
            
            batch_analysis = make_llm_api_call(client, provider, models, batch_prompt)
            
            end_time = datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            logging.info(f"Batch {batch_num + 1} successfully analyzed")
            logging.info(f"Processing time: {processing_time:.2f} seconds")
            logging.info(f"Analysis length: {len(batch_analysis)} characters")
            
            # Save individual batch analysis to file
            batch_filename = os.path.join(OUTPUT_DIRECTORY, f"batch_analysis_{batch_num + 1}.txt")
            with open(batch_filename, 'w', encoding='utf-8') as f:
                f.write(f"Batch {batch_num + 1} Analysis\n")
                f.write(f"Timestamp: {datetime.now()}\n")
                f.write(f"Pairs analyzed: {len(current_batch)}\n")
                f.write("="*50 + "\n\n")
                f.write(batch_analysis)
            
            logging.info(f"Batch analysis saved to: {batch_filename}")
            
            batch_analyses.append(batch_analysis)
            
        except Exception as e:
            logging.error(f"ERROR in batch {batch_num + 1} analysis:")
            logging.error(str(e))
            logging.error("Stack trace:", exc_info=True)
            continue

    # Final synthesis
    logging.info("\n" + "="*50)
    logging.info("STARTING FINAL SYNTHESIS")
    logging.info(f"Combining analyses from {len(batch_analyses)} successful batches")
    
    synthesis_prompt = f"""# TVÁ IDENTITA: Jsi expertním datovým analytikem a legálním poradcem, který vytváří finální a kompletní analýzu z dílčích analýz v "ZDROJOVÉ ANALÝZY" konverzací mezi uživateli a AI.

    # KONTEXT:
    Máš k dispozici {len(batch_analyses)} dílčích analýz, kde každá obsahuje extrahované páry nejčastějších dotazů a odpovědí v "ZDROJOVÉ ANALÝZY"

    # TVŮJ ULTIMÁTNÍ ÚKOL:
    1. Analyzuj všechny dílčí analýzy v "ZDROJOVÉ ANALÝZY" a identifikuj skutečně NEJČASTĚJŠÍ a NEJDŮLEŽITĚJŠÍ témata dotazů a jejich variace
    2. **SLEDUJ ČETNOSTI A SESKUPUJ INTELIGENTNĚ:** Pro každé téma agreguj celkové počty výskytů ze všech dílčích analýz a seskupuj významově totožné dotazy s drobnými variacemi (např. "dej mi číslo na Frontze" + "dej mi číslo na pana Frontze" = celkem 2x)
    3. Vyber a seřaď PŘESNĚ 15 NEJRELEVANTNĚJŠÍCH témat podle četnosti jejich výskytů
    4. **IDENTIFIKUJ CIZÍ JAZYKY:** Rozpoznej všechny dotazy v jiných jazycích než češtině a pro každý jazyk identifikuj 3 nejčastější dotazy
    5. DŮLEŽITÉ: Poskytni KOMPLETNÍ analýzu všech 15 témat najednou, bez rozdělování do částí
    6. Na konci poskytni souhrnnou statistiku
    7. **DŮLEŽITÉ - POUZE RELEVANTNÍ DOTAZY:** Zaměř se výhradně na relevantní dotazy, které vyžadovaly informace k zodpovězení

    -----------------------

    # VYŽADOVANÝ FORMÁT VÝSTUPU: Pro každé z 15 témat použij tento formát:

    ## **[POŘADOVÉ ČÍSLO]** | [SOUHRNNÝ NÁZEV TÉMATU] ([CELKOVÝ POČET VÝSKYTŮ]x)
    ### ([SEZNAM NEJČASTĚJŠÍCH VARIANT s agregovanými četnostmi - např. "kdy byla uzavřena Duchcovská?" (3x), "uzavírka Duchcovské ulice" (2x)])
    * **SUMARIZACE ODPOVĚĎÍ CHATBOTA:** [SOUHRNNÉ SUMARIZACE JAKÝM ZPŮSOBEM CHATBOT DOKÁZAL ODPOVĚDĚT - včetně kvality strukturování (URL, čísla, data, faktické informace)]

    [VLOŽIT 1 PRÁZDNÝ ŘÁDEK A POKRAČOVAT AŽ DO TÉMATU 15.]

    # SOUHRNNÁ STATISTIKA:
    
    ## CELKOVÁ PŘESNOST ANALÝZY
    * **PRŮMĚRNÁ MÍRA PŘESNOSTI:** [VYPOČÍTEJ PRŮMĚR VŠECH MĚR PŘESNOSTI Z VÝŠE UVEDENÝCH 15 TÉMAT]%
    * **ZDŮVODNĚNÍ:** [STRUČNÉ VYSVĚTLENÍ PRŮMĚRNÉ MÍRY PŘESNOSTI]

    ## CELKOVÁ ZPĚTNÁ VAZBA
    * **CELKOVÁ SPOKOJENOST UŽIVATELŮ:** [PRŮMĚR VŠECH POZITIVNÍCH ZPĚTNÝCH VAZEB]%
    * **ANALÝZA TRENDU:** [IDENTIFIKACE HLAVNÍCH DŮVODŮ SPOKOJENOSTI/NESPOKOJENOSTI]

    ## JAZYKOVÁ DISTRIBUCE: 
    * **ČEŠTINA:** [PROCENTO]%
    * **OSTATNÍ JAZYKY:** Pokud byly identifikovány dotazy v jiných jazycích, pro každý jazyk uveď:
      - **[NÁZEV JAZYKA]:** [PROCENTO]%
      - **TOP 3 NEJČASTĚJŠÍ DOTAZY:** 1) "[dotaz 1]", 2) "[dotaz 2]", 3) "[dotaz 3]"

    ## ANALÝZA NEJVYTĚŽOVANĚJŠÍCH HODIN
    * **NEJAKTIVNĚJŠÍ HODINY:** [IDENTIFIKUJ 3-5 KONKRÉTNÍCH HODIN V PRŮBĚHU DNE, KDY BYLO NEJVÍCE DOTAZŮ]
    * **ANALÝZA TRENDŮ:** [VYSVĚTLENÍ PROČ JSOU TYTO HODINY NEJVYTĚŽOVANĚJŠÍ]

    ## ODHADOVANÁ UŠETŘENÁ PRÁCA A FINANCE
    * **CELKOVÝ POČET DOTAZŮ:** {total_human_count} dotazů (ze statistik v Excel reportu)
    * **ODHADOVANÉ UŠETŘENÉ HODINY:** {total_human_count * 0.25} hodin (předpoklad 15 minut na dotaz)
    * **ODHADOVANÉ UŠETŘENÉ FINANCE:** {total_human_count * 0.25 * 200} Kč (při sazbě 200 Kč/hodina pro úředníka)
    * **ROČNÍ PROJEKCE:** [EXTRAPOLACE NA CELÝ ROK PODLE SOUČASNÉHO TRENDU]
    -----------------------
        
    # CELKOVÁ OMEZENÍ A POŽADAVKY:
    - Musíš analyzovat a vrátit VŠECH 15 témat najednou
    - **POVINNÉ ČETNOSTI:** Pro každé téma a variantu MUSÍŠ uvést agregované počty výskytů ze všech dílčích analýz ve formátu "(Xx)"
    - **INTELIGENTNÍ SESKUPOVÁNÍ:** Seskupuj významově totožné dotazy i s drobnými variacemi (např. "dej mi číslo na Frontze" + "dej mi číslo na pana Frontze" = "dej mi číslo na (pana) Frontze" (2x))
    - **POVINNÉ CIZÍ JAZYKY:** Pokud existují dotazy v jiných jazycích než češtině, MUSÍŠ je identifikovat a uvést top 3 pro každý jazyk (BEZ ČÍSEL VÝSKYTŮ - POUZE PROCENTA)
    - **POUZE RELEVANTNÍ DOTAZY:** Zaměř se výhradně na relevantní dotazy, které vyžadovaly relevantní informace k zodpovězení
    - Nepoužívej žádné poznámky o pokračování nebo rozdělování odpovědi
    - Drž se striktně daného formátu pro každé téma
    - Vynechej na začátku tvého výstupu jakékoliv poznámky, či sebe vysvětlování => ihned začni s generováním jednotlivých dotazů a odpovědí
    - Pro každé téma MUSÍŠ analyzovat zpětnou vazbu uživatelů
    - Na konci MUSÍŠ uvést souhrnnou statistiku s průměrnou mírou přesnosti a celkovou zpětnou vazbou
    - **POVINNÉ NOVÉ SEKCE:** MUSÍŠ přidat analýzu nejvytěžovanějších hodin a odhadovanou ušetřenou práci a finance

    # "ZDROJOVÉ ANALÝZY":
    {chr(10).join([f"=== BATCH {i+1} ===\n{analysis}\n" for i, analysis in enumerate(batch_analyses)])}
    """

    try:
        logging.info("Making final synthesis API call...")
        start_time = datetime.now()
        
        final_analysis = make_llm_api_call(client, provider, models, synthesis_prompt)
        
        end_time = datetime.now()
        total_processing_time = (end_time - start_time).total_seconds()
        
        logging.info("Final synthesis completed successfully")
        logging.info(f"Total processing time: {total_processing_time:.2f} seconds")
        logging.info(f"Final analysis length: {len(final_analysis)} characters")
        
        # Save final synthesis to separate file
        synthesis_filename = os.path.join(OUTPUT_DIRECTORY, "final_synthesis.txt")
        with open(synthesis_filename, 'w', encoding='utf-8') as f:
            f.write(f"Final Synthesis Report\n")
            f.write(f"Generated: {datetime.now()}\n")
            f.write(f"Total batches analyzed: {total_batches}\n")
            f.write(f"Total Q&A pairs extracted: {len(qa_pairs)}\n")
            f.write(f"Total human messages (Excel): {total_human_count}\n")
            f.write("="*50 + "\n\n")
            f.write(final_analysis)
        
        logging.info(f"Final synthesis saved to: {synthesis_filename}")
        logging.info("="*80)
        
        return final_analysis
        
    except Exception as e:
        logging.error("ERROR in final synthesis:")
        logging.error(str(e))
        logging.error("Stack trace:", exc_info=True)
        raise

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Chyba: Nebyl zadán název konfiguračního souboru.")
        print("Použití: python ExportConvos_v2.py <název_konfiguračního_souboru>")
        sys.exit(1)
    
    config_file = sys.argv[1]
    main(config_file)