import sys
import os
import json
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
import re
import matplotlib.pyplot as plt
from io import BytesIO
from openpyxl.drawing.image import Image
import numpy as np
import logging
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import requests

# Funkce pro načtení konfigurace ze souboru
def load_config(config_file):
    logging.info(f"Loading configuration from {config_file}")
    config = {}
    try:
        with open(config_file, 'r') as file:
            for line in file:
                key, value = line.strip().split('=', 1)
                config[key] = value.strip()
                
        # Přidat další potřebné konfigurační hodnoty
        config['START_DATE'] = config.get('START_DATE', '')
        config['END_DATE'] = config.get('END_DATE', '')
        config['CATEGORIES'] = config.get('CATEGORIES', '').split(',')
        
        logging.info("Configuration loaded successfully")
        logging.debug(f"Config contents: {config}")
        return config
    except Exception as e:
        logging.error(f"Error loading config: {str(e)}")
        raise

# Set up logging configuration
def setup_logging(output_directory):
    log_filepath = os.path.join(output_directory, 'execution.log')
    
    # First remove any existing handlers to avoid duplicates
    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    
    # Configure logging with immediate output
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    
    file_handler = logging.FileHandler(log_filepath, encoding='utf-8', mode='w')
    file_handler.setLevel(logging.INFO)
    
    # Create formatter and add it to the handlers
    formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
    console_handler.setFormatter(formatter)
    file_handler.setFormatter(formatter)
    
    # Get the root logger and add handlers
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    logger.addHandler(console_handler)
    logger.addHandler(file_handler)
    
    # Force immediate flush
    console_handler.flush()
    file_handler.flush()
    
    # Test logging is working
    logging.info("="*80)
    logging.info("LOGGING INITIALIZED SUCCESSFULLY")
    logging.info(f"Log file created at: {log_filepath}")
    logging.info("="*80)
    
    return log_filepath

def main(config_file):
    try:
        config = load_config(config_file)
        
        global PROJECT_NAME, OUTPUT_DIRECTORY, ANTHROPIC_API_KEY, EXCEL_FILE, START_DATE, END_DATE, CATEGORIES
        PROJECT_NAME = config['PROJECT_NAME']
        ANTHROPIC_API_KEY = config['ANTHROPIC_API_KEY']
        EXCEL_FILE = config['EXCEL_FILE']
        START_DATE = config['START_DATE']
        END_DATE = config['END_DATE']
        CATEGORIES = config['CATEGORIES']
        
        # Create output directory name from Excel filename without extension
        excel_name = os.path.splitext(os.path.basename(EXCEL_FILE))[0]
        OUTPUT_DIRECTORY = f"{PROJECT_NAME}_{excel_name}"
        
        # Create directory and set up logging
        create_output_directory()
        setup_logging(OUTPUT_DIRECTORY)
        
        # Read Excel file
        df = pd.read_excel(EXCEL_FILE)
        
        # Analyze conversations with Claude
        logging.info("Analyzing conversations with Claude...")
        faq_analysis = analyze_excel_data_with_claude(df)
        save_faq_to_word(faq_analysis, excel_name)
        
        logging.info("Script execution completed successfully")
    except Exception as e:
        logging.error(f"Critical error in main execution: {str(e)}", exc_info=True)
        raise

def create_output_directory():
    logging.info(f"Creating output directory: {OUTPUT_DIRECTORY}")
    try:
        if not os.path.exists(OUTPUT_DIRECTORY):
            os.makedirs(OUTPUT_DIRECTORY)
            logging.info(f"Created directory: {OUTPUT_DIRECTORY}")
        else:
            logging.info(f"Directory already exists: {OUTPUT_DIRECTORY}")
    except Exception as e:
        logging.error(f"Error creating output directory: {str(e)}")
        raise

def get_transcript_ids(base_url, headers, project_id):
    logging.info(f"Fetching transcript IDs for project {project_id}")
    try:
        url = f"{base_url}/{project_id}"
        params = {
            "startDate": START_DATE,
            "endDate": END_DATE
        }
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        transcript_ids = [transcript["_id"] for transcript in response.json()]
        logging.info(f"Retrieved {len(transcript_ids)} transcript IDs")
        return transcript_ids
    except Exception as e:
        logging.error(f"Error fetching transcript IDs: {str(e)}")
        raise

def get_transcript_dialog(base_url, headers, project_id, transcript_id):
    logging.info(f"Fetching dialog for transcript {transcript_id}")
    try:
        url = f"{base_url}/{project_id}/{transcript_id}"
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        dialog = response.json()
        logging.info(f"Successfully retrieved dialog for transcript {transcript_id} - {len(dialog)} messages")
        return dialog
    except Exception as e:
        logging.error(f"Error fetching transcript dialog {transcript_id}: {str(e)}")
        raise

def extract_messages(dialog):
    messages = []
    for turn in dialog:
        if turn["type"] == "debug" and "payload" in turn and "payload" in turn["payload"]:
            debug_message = turn["payload"]["payload"].get("message", "")
            if "CategoryFilter" in debug_message:
                messages.append({
                    "role": "DEBUG",
                    "content": debug_message,
                    "timestamp": turn.get("startTime", "")
                })
        elif turn["type"] == "request":
            if "payload" in turn and "query" in turn["payload"].get("payload", {}):
                messages.append({
                    "role": "HUMAN",
                    "content": turn["payload"]["payload"]["query"],
                    "timestamp": turn.get("startTime", "")
                })
        elif turn["type"] == "text" and "payload" in turn and "message" in turn["payload"].get("payload", {}):
            messages.append({
                "role": "BOT",
                "content": turn["payload"]["payload"]["message"],
                "timestamp": turn.get("startTime", "")
            })
    return messages

def save_transcript_to_txt(transcript_id, messages):
    filename = os.path.join(OUTPUT_DIRECTORY, f"transcript_{transcript_id}.txt")
    logging.info(f"Saving transcript {transcript_id} to file...")
    sys.stdout.flush()  # Force immediate console output
    
    try:
        message_count = len(messages)
        with open(filename, 'w', encoding='utf-8') as file:
            for idx, message in enumerate(messages, 1):
                if message['role'] == 'DEBUG':
                    file.write(f"DEBUG: {message['content']}\n")
                else:
                    file.write(f"{message['role']}: {message['content']}\n")
                file.write("----------\n")
        logging.info(f"✓ Successfully saved transcript {transcript_id} ({message_count} messages)")
        sys.stdout.flush()  # Force immediate console output
    except Exception as e:
        logging.error(f"Error saving transcript {transcript_id}: {str(e)}")
        sys.stdout.flush()  # Force immediate console output
        raise

def count_human_occurrences():
    human_count = 0
    for filename in os.listdir(OUTPUT_DIRECTORY):
        if filename.startswith("transcript_") and filename.endswith(".txt"):
            file_path = os.path.join(OUTPUT_DIRECTORY, filename)
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                human_count += content.count("HUMAN:")
    return human_count

def count_category_occurrences():
    category_counts = {category: 0 for category in CATEGORIES}
    
    for filename in os.listdir(OUTPUT_DIRECTORY):
        if filename.startswith("transcript_") and filename.endswith(".txt"):
            file_path = os.path.join(OUTPUT_DIRECTORY, filename)
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                for category in CATEGORIES:
                    exact_match = f'\\"{category}\\"'
                    category_counts[category] += content.count(exact_match)
    
    return category_counts

def create_pie_chart(category_counts):
    labels = list(category_counts.keys())
    sizes = list(category_counts.values())
    
    plt.figure(figsize=(10, 8))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    plt.axis('equal')
    plt.title('Rozložení kategorií')
    
    # Uložení grafu jako obrázek
    img_path = os.path.join(OUTPUT_DIRECTORY, 'category_distribution.png')
    plt.savefig(img_path)
    plt.close()
    
    return img_path

def create_custom_donut_chart(category_counts):
    logging.info("Creating custom donut chart")
    try:
        # Seřazení kategorií sestupně podle počtu
        sorted_categories = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)
        labels = [cat for cat, count in sorted_categories if count > 0]
        sizes = [count for _, count in sorted_categories if count > 0]

        # Vytvoření širšího spektra barev duhy v opačném pořadí
        num_colors = len(labels)
        rainbow_colors = plt.cm.rainbow(np.linspace(1, 0, num_colors))

        # Změna velikosti figury na čtvercový tvar
        fig, ax = plt.subplots(figsize=(20, 20), subplot_kw=dict(aspect="equal"))  

        wedges, texts, autotexts = ax.pie(sizes, wedgeprops=dict(width=0.5), startangle=-40,
                                          colors=rainbow_colors, autopct='%1.1f%%', pctdistance=0.85)

        bbox_props = dict(boxstyle="round,pad=0.3", fc="w", ec="k", lw=0.72)
        kw = dict(arrowprops=dict(arrowstyle="-", connectionstyle="angle,angleA=0,angleB=90,rad=10"),
                  bbox=bbox_props, zorder=0, va="center")

        # Úprava umístění popisků
        for i, p in enumerate(wedges):
            ang = (p.theta2 - p.theta1) / 2. + p.theta1
            y = np.sin(np.deg2rad(ang))
            x = np.cos(np.deg2rad(ang))
            horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
            connectionstyle = f"angle,angleA=0,angleB={ang}"
            kw["arrowprops"].update({"connectionstyle": connectionstyle})
            ax.annotate(labels[i], xy=(x, y), xytext=(2.2*np.sign(x), 2.2*y),
                        horizontalalignment=horizontalalignment, fontsize=14, **kw)

        plt.title(f"Rozložení kategorií\n{START_DATE} - {END_DATE}", fontsize=24, y=1.05)

        for autotext in autotexts:
            autotext.set_visible(False)

        # Odstranění os pro čistší vzhled
        ax.set_axis_off()

        # Uložení grafu
        img_path = os.path.join(OUTPUT_DIRECTORY, 'custom_category_distribution.png')
        plt.savefig(img_path, bbox_inches='tight', dpi=300)
        plt.close()
        logging.info(f"Donut chart saved to {img_path}")
        return img_path
    except Exception as e:
        logging.error(f"Error creating donut chart: {str(e)}")
        raise

def print_summary(human_count, category_counts):
    print("\nSouhrn zpracování:")
    print(f"1) Počet HUMAN výskytů: {human_count}")
    print("\n2) Počty výskytů kategorií:")
    for category, count in category_counts.items():
        print(f"   {category}: {count}")

def analyze_excel_data_with_claude(df):
    """Analyze Excel data using Claude"""
    logging.info("\n" + "="*80)
    logging.info("ZAČÁTEK ANALÝZY DAT POMOCÍ CLAUDE")
    logging.info("="*80)
    
    client = Anthropic(api_key=ANTHROPIC_API_KEY)
    
    # Get Excel filename without extension for project name
    excel_name = os.path.splitext(os.path.basename(EXCEL_FILE))[0]
    project_name = f'Analýza zpětné vazby {excel_name}'
    
    # Prepare data for analysis
    conversations = []
    logging.info("\nZpracování Excel dat:")
    logging.info("-"*50)
    
    for index, row in df.iterrows():
        logging.info(f"\nZpracovávám řádek {index + 1}/{len(df)}")
        poznamka = str(row['Poznámka']) if pd.notna(row['Poznámka']) else ''
        messages = str(row['Zprávy']) if pd.notna(row['Zprávy']) else ''
        
        if poznamka:
            conversations.append({
                'Poznámka': poznamka,
                'Zprávy': messages,
                'has_feedback': bool(poznamka.strip())
            })
            logging.info(f"✓ Řádek {index + 1} obsahuje zpětnou vazbu a byl přidán do analýzy")
            logging.info(f"  • Délka poznámky (zpětná vazba): {len(poznamka)} znaků")
            logging.info(f"  • Délka zpráv: {len(messages)} znaků")
        else:
            logging.info(f"⚠ Řádek {index + 1} přeskočen - neobsahuje zpětnou vazbu v poznámce")

    def create_analysis_prompt(data_to_analyze):
        """Vytvoří jednotný prompt pro analýzu dat"""
        return f"""# KRITICKÁ PRAVIDLA:
        - Nejprve SESKUP všechny konverzace podle TÉMATICKY PODOBNÉHO znění zpětné vazby v poli 'Poznámka'
        - Pro KAŽDOU unikátní zpětnou vazbu:
          1. Spočítej její četnost výskytu
          2. Seřaď všechny skupiny sestupně podle četnosti
        - Pro KAŽDOU skupinu vypiš:
          1. Četnost výskytu daných zpětných vazeb
          2. PŘESNÉ dílčí znění jednotlivých zpětných vazeb
          3. VŠECHNY příslušné konverzace s touto zpětnou vazbou
        - Pokračuj, dokud nejsou vypsány VŠECHNY existující zpětné vazby
        - NEZAHRNUJ žádné analýzy ani interpretace

        # POŽADOVANÝ FORMÁT VÝSTUPU:

        ## SKUPINA 1 (výskytů: X)
        **Zpětná vazba:**
        [PŘESNÁ citace z pole Poznámka]

        **Konverzace 1 z X:**
        Osoba: [PŘESNÝ dotaz]
        Asistent: [PŘESNÁ odpověď]
        [všechny zprávy v konverzaci]

        **Konverzace 2 z X:**
        [stejný formát pro všechny konverzace ve skupině]

        ## SKUPINA 2 (výskytů: Y)
        [stejný formát]

        [Pokračuj pro VŠECHNY existující skupiny]

        # VSTUPNÍ DATA K ANALÝZE:
        {json.dumps(data_to_analyze, ensure_ascii=False, indent=2)}
        """

    # Batch processing
    batch_size = 10
    total_batches = (len(conversations) + batch_size - 1) // batch_size
    batch_analyses = []
    all_analyzed_data = []  # Pro ukládání všech analyzovaných dat
    
    logging.info("\nZahajuji zpracování v dávkách:")
    logging.info(f"• Velikost dávky: {batch_size} konverzací")
    logging.info(f"• Celkový počet dávek: {total_batches}")
    
    for batch_num in range(total_batches):
        start_idx = batch_num * batch_size
        end_idx = min((batch_num + 1) * batch_size, len(conversations))
        current_batch = conversations[start_idx:end_idx]
        
        logging.info(f"\nZpracovávám dávku {batch_num + 1}/{total_batches}")
        logging.info(f"• Konverzace {start_idx + 1} až {end_idx}")
        
        try:
            batch_prompt = create_analysis_prompt(current_batch)
            
            logging.info("Odesílám požadavek na analýzu dávky...")
            message = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=8000,  # Snížený limit tokenů na maximum povolené hodnoty
                temperature=0,
                messages=[
                    {"role": "user", "content": batch_prompt}
                ]
            )
            
            batch_analysis = message.content[0].text.strip()
            batch_analyses.append(batch_analysis)
            all_analyzed_data.extend(current_batch)  # Přidáme data do celkové analýzy
            
            # Uložení dílčí analýzy do samostatného souboru
            batch_filename = os.path.join(OUTPUT_DIRECTORY, f'batch_analysis_{batch_num + 1}.txt')
            with open(batch_filename, 'w', encoding='utf-8') as f:
                f.write(f"=== ANALÝZA DÁVKY {batch_num + 1}/{total_batches} ===\n\n")
                f.write(f"Konverzace {start_idx + 1} až {end_idx}\n")
                f.write("="*50 + "\n\n")
                f.write(batch_analysis)
            
            logging.info(f"✓ Dávka {batch_num + 1} úspěšně analyzována a uložena")
            logging.info(f"  • Délka analýzy: {len(batch_analysis)} znaků")
            logging.info(f"  • Uloženo do: {batch_filename}")
            
        except Exception as e:
            logging.error(f"❌ Chyba při zpracování dávky {batch_num + 1}: {str(e)}")
            continue

    # Finální syntéza používá všechna sesbíraná data
    try:
        logging.info("Odesílám požadavek na finální analýzu...")
        final_prompt = create_analysis_prompt(all_analyzed_data)  # Použijeme sesbíraná data
        
        message = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=8000,  # Snížený limit tokenů
            temperature=0,
            messages=[
                {"role": "user", "content": final_prompt}
            ]
        )
        
        final_analysis = message.content[0].text.strip()
        
        # Save the final analysis
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        analysis_file = os.path.join(OUTPUT_DIRECTORY, f'final_analysis_{timestamp}.txt')
        with open(analysis_file, 'w', encoding='utf-8') as f:
            f.write(final_analysis)
            
        logging.info(f"✓ Finální analýza uložena do: {analysis_file}")
        logging.info(f"  • Délka analýzy: {len(final_analysis)} znaků")
        
        return final_analysis
        
    except Exception as e:
        logging.error(f"❌ Chyba při vytváření finální analýzy: {str(e)}")
        raise

def save_faq_to_word(analysis_text, excel_name):
    doc = Document()
    
    # Nastavení stylu nadpisu
    title = doc.add_heading('', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title text with bold excel filename
    title_text = title.add_run(f"ANALÝZA ZPĚTNÉ VAZBY ")
    filename = title.add_run(excel_name)
    filename.bold = True
    
    doc.add_paragraph()
    
    # Rozdělení analýzy na řádky
    lines = analysis_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Zpracování markdown nadpisů
        if line.startswith('#'):
            # Určení úrovně nadpisu podle počtu #
            level = len(line) - len(line.lstrip('#'))
            # Odstranění # a případného | ze začátku textu
            heading_text = line.lstrip('#').strip()
            if '|' in heading_text:
                heading_text = heading_text.split('|', 1)[1].strip()
            
            # Přidání nadpisu odpovídající úrovně (0-8)
            heading = doc.add_heading(heading_text, level)
            continue
        
        # Zpracování odrážek (bullet points)
        if line.startswith('*'):
            text = line[1:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            
            # Zpracování tučného textu mezi **
            if '**' in text:
                p.clear()
                parts = text.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Lichá část (mezi ** **)
                        p.add_run(part).bold = True
                    else:  # Sudá část (normální text)
                        p.add_run(part)
        
        # Běžný text
        else:
            p = doc.add_paragraph(line)
    
    # Nastavení stylu dokumentu
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(12)
        for run in paragraph.runs:
            run.font.size = Pt(11)
    
    # Uložení dokumentu - using the same name as Excel file
    output_filename = f'{excel_name}.docx'
    output_path = os.path.join(OUTPUT_DIRECTORY, output_filename)
    doc.save(output_path)
    logging.info(f"FAQ analýza uložena do Word dokumentu: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Chyba: Nebyl zadán název konfiguračního souboru.")
        print("Použití: python ExportConvos_v2.py <název_konfiguračního_souboru>")
        sys.exit(1)
    
    config_file = sys.argv[1]
    main(config_file)